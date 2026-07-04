"""HR/legal document generation: LLM drafts the text (grounded in legislation
via RAG), python-docx renders a downloadable DOCX.

The text→DOCX conversion is a pure function (`text_to_docx_paragraphs` +
`build_docx`) so it is unit-testable without an LLM.
"""

import io
import uuid
from dataclasses import dataclass

from sqlalchemy.ext.asyncio import AsyncSession

from app.services.ai.base import ChatMessage
from app.services.ai.registry import get_provider
from app.services.rag.pipeline import retrieve
from app.services.security.guard import screen_user_input, wrap_retrieved_context

DOC_TYPES: dict[str, str] = {
    "order": "Приказ по кадрам (о приеме, увольнении, отпуске, командировке и т.п.)",
    "vacation_request": "Заявление на отпуск",
    "explanation_letter": "Объяснительная записка",
    "employment_contract": "Трудовой договор",
    "dismissal_notice": "Уведомление о прекращении трудового договора",
    "reference": "Справка с места работы",
    "other": "Иной кадровый или юридический документ",
}

_SYSTEM_PROMPT = """Ты — эксперт по кадровому делопроизводству Республики Узбекистан.
Составь готовый к использованию документ по запросу пользователя.

Требования:
- строго соответствуй Трудовому кодексу РУз; где уместно, ссылайся на статьи;
- используй официально-деловой стиль и стандартные реквизиты (шапка, дата, подпись);
- недостающие данные обозначай плейсхолдерами вида [ФИО работника], [дата];
- отвечай ТОЛЬКО текстом документа, без пояснений до или после;
- строки, начинающиеся с '# ', считаются заголовком документа, с '## ' — подзаголовком;
- пиши на языке запроса пользователя (узбекский или русский)."""


@dataclass
class GeneratedDocument:
    title: str
    text: str
    docx: bytes
    sources: list[dict]


def text_to_docx_paragraphs(text: str) -> list[tuple[str, str]]:
    """Parse generated text into (style, content) pairs: title/heading/body."""
    paragraphs: list[tuple[str, str]] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("## "):
            paragraphs.append(("heading", line[3:].strip()))
        elif line.startswith("# "):
            paragraphs.append(("title", line[2:].strip()))
        else:
            paragraphs.append(("body", line.strip()))
    return paragraphs


def build_docx(text: str) -> bytes:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    docx = DocxDocument()
    for style, content in text_to_docx_paragraphs(text):
        if style == "title":
            p = docx.add_heading(content, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style == "heading":
            docx.add_heading(content, level=2)
        else:
            docx.add_paragraph(content)
    buffer = io.BytesIO()
    docx.save(buffer)
    return buffer.getvalue()


async def generate_document(
    db: AsyncSession,
    tenant_id: uuid.UUID,
    *,
    doc_type: str,
    instructions: str,
    provider_name: str | None = None,
) -> GeneratedDocument:
    verdict = screen_user_input(instructions)
    if not verdict.allowed:
        raise ValueError("Instructions rejected by the security screen")
    type_hint = DOC_TYPES.get(doc_type, DOC_TYPES["other"])

    rag = await retrieve(db, tenant_id, instructions, top_k=4, use_reranker=False)
    messages = [ChatMessage(role="system", content=_SYSTEM_PROMPT)]
    if rag.context:
        messages.append(ChatMessage(role="system", content=wrap_retrieved_context(rag.context)))
    messages.append(
        ChatMessage(role="user", content=f"Тип документа: {type_hint}.\n\nДанные для документа:\n{instructions}")
    )

    result = await get_provider(provider_name).complete(messages, temperature=0.1, max_tokens=3000)
    text = result.content.strip()
    parsed = text_to_docx_paragraphs(text)
    title = next((c for s, c in parsed if s == "title"), type_hint)
    return GeneratedDocument(title=title, text=text, docx=build_docx(text), sources=rag.sources)
